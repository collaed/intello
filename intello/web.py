"""FastAPI web interface for L'Intello."""
import asyncio
import json
import os
import re
import shutil
import time
import uuid
from typing import Optional

import base64
from fastapi import FastAPI, UploadFile, File, Form, Request
from fastapi.responses import HTMLResponse, RedirectResponse, JSONResponse, Response, StreamingResponse
from starlette.middleware.base import BaseHTTPMiddleware

from intello.models import Tier
from intello.research import get_providers, probe_reference_sites
from intello.keys import discover_keys, validate_keys, add_key
from intello.router import build_plan
from intello.backends import execute, SYSTEM_DEFAULT
from intello import gdrive
from intello import ratelimit
from intello.pipeline import run_deep
from intello import memory
from intello import cache
from intello.chains import analyze_complexity, execute_chain
from intello.tools import TOOL_PROMPT_SUFFIX, detect_tool_call, execute_tool
from intello.guardrails import check_confidence
from intello.debate import run_debate
from intello import literary
from intello.craft import build_craft_prompt
from intello.guardrails import check_confidence, check_word_count
from intello import workflow as wf
from intello import writing_tools as wt
from intello import ocr
from intello import ocr_engines
from intello import scheduler
from intello import imagegen
from intello import webhooks
from intello import reconstruct as recon

app = FastAPI(title="L'Intello")

# Auth — all credentials from environment variables
import json as _json
_users_raw = os.environ.get("INTELLO_USERS", '{"admin": "changeme"}')
try:
    USERS = _json.loads(_users_raw)
except Exception:
    USERS = {"admin": "changeme"}
TOKEN = os.environ.get("INTELLO_TOKEN", "changeme")
PREMIUM_USERS = set(os.environ.get("INTELLO_PREMIUM_USERS", "admin").split(","))

# Models restricted to specific users (everyone else gets them filtered out)
PREMIUM_MODELS = {
    "gemini-2.5-pro",
    "claude-sonnet-4-5",
    "gpt-4o",
    "grok-4-1-fast",
}
# PREMIUM_USERS is set above from env var


def _get_user(request: Request) -> str:
    """Extract current user from request."""
    # From Caddy forward_auth
    user = request.headers.get("X-Auth-User", "")
    if user:
        return user
    # From Basic auth
    auth = request.headers.get("Authorization", "")
    if auth.startswith("Basic "):
        try:
            import base64
            decoded = base64.b64decode(auth[6:]).decode()
            return decoded.split(":", 1)[0]
        except Exception:
            pass
    # Docker internal = admin
    client_ip = request.client.host if request.client else ""
    if client_ip.startswith("172.") or client_ip == "127.0.0.1":
        return "ecb"
    return "anonymous"


def filter_providers_for_user(providers: list, user: str) -> list:
    """Filter out premium models for non-premium users."""
    if user in PREMIUM_USERS:
        return providers
    return [p for p in providers if not any(pm in p.model_id for pm in PREMIUM_MODELS)]


class AuthMiddleware(BaseHTTPMiddleware):
    async def dispatch(self, request: Request, call_next):
        # Trust Docker internal network (other containers)
        client_ip = request.client.host if request.client else ""
        if client_ip.startswith("172.") or client_ip == "127.0.0.1":
            return await call_next(request)
        # Caddy forward_auth sets this header — trust it
        if request.headers.get("X-Auth-User"):
            return await call_next(request)
        auth = request.headers.get("Authorization", "")
        # Bearer token (for API clients like audiobookshelf)
        if auth.startswith("Bearer "):
            if auth[7:] == TOKEN:
                return await call_next(request)
        # Basic auth
        if auth.startswith("Basic "):
            try:
                decoded = base64.b64decode(auth[6:]).decode()
                user, pwd = decoded.split(":", 1)
                if USERS.get(user) == pwd:
                    return await call_next(request)
            except Exception:
                pass
        # Cookie
        if request.cookies.get("intello_token") == TOKEN:
            return await call_next(request)
        # Query param
        if request.query_params.get("token") == TOKEN:
            return await call_next(request)
        # Login endpoint
        if request.url.path == "/login":
            return await call_next(request)
        # Login page for direct access (no auth proxy)
        if request.url.path in ("/", "/literary") and request.method == "GET":
            return HTMLResponse(_login_page())
        return Response("Unauthorized", status_code=401,
                        headers={"WWW-Authenticate": 'Basic realm="Intello"'})


app.add_middleware(AuthMiddleware)


@app.post("/login")
async def login(request: Request, password: str = Form(...)):
    if password == TOKEN:
        # Detect base path from request
        base = request.headers.get("X-Forwarded-Prefix", "")
        resp = RedirectResponse(base + "/", status_code=303)
        resp.set_cookie("intello_token", TOKEN, httponly=True, max_age=86400 * 30,
                        samesite="lax", path="/")
        return resp
    return HTMLResponse(_login_page("Wrong password"))


def _login_page(error=""):
    return f"""<!DOCTYPE html><html><head><title>L&#39;Intello Login</title>
<style>body{{background:#0f1117;color:#e4e4e7;font-family:sans-serif;display:flex;align-items:center;justify-content:center;height:100vh;margin:0}}
.box{{background:#1a1d27;padding:32px;border-radius:12px;border:1px solid #2a2d3a;text-align:center}}
input{{background:#0f1117;border:1px solid #2a2d3a;color:#e4e4e7;padding:10px;border-radius:6px;margin:8px 0;font-size:1rem}}
button{{background:#6366f1;color:#fff;border:none;padding:10px 24px;border-radius:6px;cursor:pointer;font-size:1rem}}
.err{{color:#ef4444;font-size:.85rem}}</style></head>
<body><div class="box"><h2>⚡ L'Intello</h2>
<form method="POST" action="login"><input type="password" name="password" placeholder="Password" autofocus>
<br><button type="submit">Login</button></form>
{f'<p class="err">{error}</p>' if error else ''}</div></body></html>"""

_providers = []


@app.on_event("startup")
async def startup():
    global _providers
    _providers = get_providers()
    discover_keys(_providers)
    await asyncio.gather(
        probe_reference_sites(),
        validate_keys(_providers),
    )
    # Preload sentence-transformer in background (avoids first-request timeout)
    asyncio.get_event_loop().run_in_executor(None, cache._embedder)


def _provider_dict(p):
    rem = ratelimit.remaining(p.model_id, p.daily_limit)
    return {
        "name": p.name, "model_id": p.model_id, "provider": p.provider,
        "tier": p.tier.value, "available": p.available,
        "env_key": p.env_key, "has_key": p.api_key is not None,
        "cost_per_1k_input": p.cost_per_1k_input,
        "cost_per_1k_output": p.cost_per_1k_output,
        "notes": p.notes,
        "daily_limit": p.daily_limit,
        "used_today": ratelimit.get_usage(p.model_id),
        "remaining": rem,
    }


@app.get("/api/providers")
async def api_providers():
    return [_provider_dict(p) for p in _providers]


@app.post("/api/key")
async def api_add_key(env_key: str = Form(...), value: str = Form(...)):
    add_key(_providers, env_key, value)
    await validate_keys(_providers)
    return {"ok": True}


# --- Google Drive OAuth ---

@app.get("/api/gdrive/status")
async def gdrive_status():
    return {"authenticated": gdrive.is_authenticated(),
            "configured": os.path.exists(gdrive.CREDENTIALS_PATH)}


@app.get("/api/gdrive/auth")
async def gdrive_auth(request: Request):
    redirect_uri = str(request.url_for("gdrive_callback"))
    url = gdrive.get_oauth_url(redirect_uri)
    if not url:
        return {"error": "Google Drive OAuth not configured. Place gdrive_credentials.json in /data/"}
    return RedirectResponse(url)


@app.get("/api/gdrive/callback")
async def gdrive_callback(request: Request, code: str):
    redirect_uri = str(request.url_for("gdrive_callback"))
    gdrive.exchange_code(code, redirect_uri)
    return RedirectResponse("/")


# --- Google Drive Browser ---

@app.get("/api/gdrive/browse")
async def api_gdrive_browse(folder_id: str = "root", q: str = ""):
    if not gdrive.is_authenticated():
        return {"error": "Not authenticated with Google Drive"}
    return gdrive.list_folder(folder_id, q)


@app.post("/api/gdrive/batch")
async def api_gdrive_batch(request: Request):
    """Fetch multiple files by ID. Body: {"file_ids": ["id1", "id2", ...]}"""
    body = await request.json()
    file_ids = body.get("file_ids", [])
    if not file_ids:
        return {"error": "No file_ids provided"}
    return gdrive.batch_fetch(file_ids)


@app.post("/api/reconstruct/{project_id}/ingest-gdrive")
async def api_recon_ingest_gdrive(project_id: str, request: Request):
    """Ingest multiple Google Drive files into a reconstruction project."""
    body = await request.json()
    file_ids = body.get("file_ids", [])
    if not file_ids:
        return {"error": "No file_ids"}

    files = gdrive.batch_fetch(file_ids)
    results = []
    for f in files:
        if f.get("error"):
            results.append({"name": f.get("name", "?"), "error": f["error"]})
            continue
        r = recon.ingest_version(project_id, f["name"], f["content"])
        results.append({"name": f["name"], **r})

    return {"ingested": len([r for r in results if "error" not in r]),
            "errors": len([r for r in results if "error" in r]),
            "results": results}


# --- Conversations & Memory ---

@app.get("/api/conversations")
async def api_conversations():
    return memory.list_conversations()


@app.get("/api/conversations/{conv_id}")
async def api_conversation(conv_id: str):
    return {"messages": memory.get_messages(conv_id, limit=50),
            "summary": memory.get_summary(conv_id)}


# --- User Preferences ---

@app.get("/api/prefs")
async def api_get_prefs():
    return memory.get_prefs()


@app.post("/api/prefs")
async def api_set_prefs(
    tone: Optional[str] = Form(None),
    default_mode: Optional[str] = Form(None),
    custom_system_prompt: Optional[str] = Form(None),
):
    kwargs = {}
    if tone is not None: kwargs["tone"] = tone
    if default_mode is not None: kwargs["default_mode"] = default_mode
    if custom_system_prompt is not None: kwargs["custom_system_prompt"] = custom_system_prompt
    memory.set_prefs(**kwargs)
    return {"ok": True}


# --- Feedback / Learning ---

@app.post("/api/feedback")
async def api_feedback(
    model_id: str = Form(...),
    task_type: str = Form(...),
    rating: int = Form(...),
):
    memory.record_model_result(model_id, task_type, success=True, rating=float(rating))
    return {"ok": True}


@app.get("/api/learning")
async def api_learning():
    return memory.get_model_scores()


@app.get("/api/cache/stats")
async def api_cache_stats():
    return cache.get_stats()


# --- Literary Engine ---

@app.get("/api/literary/projects")
async def api_literary_projects():
    return literary.list_projects()


@app.post("/api/literary/projects")
async def api_literary_create_project(
    title: str = Form(...),
    genre: str = Form("fiction"),
    brief: str = Form(""),
    target_words: int = Form(0),
    style: str = Form(""),
    steps: str = Form("[]"),
):
    import uuid as _uuid
    pid = title.replace(" ", "_").lower()[:30] + f"_{int(time.time())}"
    try:
        steps_list = json.loads(steps) if steps else []
    except json.JSONDecodeError:
        steps_list = [s.strip() for s in steps.split("\n") if s.strip()]
    return literary.create_project(pid, title, genre, brief, target_words, style, steps_list)


@app.get("/api/literary/projects/{project_id}")
async def api_literary_get_project(project_id: str):
    p = literary.get_project(project_id)
    if not p:
        return {"error": "Project not found"}
    return p


@app.post("/api/literary/projects/{project_id}")
async def api_literary_update_project(
    project_id: str,
    title: Optional[str] = Form(None),
    genre: Optional[str] = Form(None),
    brief: Optional[str] = Form(None),
    target_words: Optional[int] = Form(None),
    style: Optional[str] = Form(None),
    steps: Optional[str] = Form(None),
):
    kwargs = {}
    if title is not None: kwargs["title"] = title
    if genre is not None: kwargs["genre"] = genre
    if brief is not None: kwargs["brief"] = brief
    if target_words is not None: kwargs["target_words"] = target_words
    if style is not None: kwargs["style"] = style
    if steps is not None:
        try:
            kwargs["steps"] = json.loads(steps)
        except json.JSONDecodeError:
            kwargs["steps"] = [s.strip() for s in steps.split("\n") if s.strip()]
    return literary.update_project(project_id, **kwargs)


@app.post("/api/literary/projects/{project_id}/auto-populate")
async def api_literary_auto_populate(project_id: str, doc_id: str = Form(...)):
    """Use LLMs to auto-fill project fields from the document text."""
    info = literary.get_document_info(doc_id)
    if not info:
        return {"error": "Document not found"}

    chunks = literary.get_chunks(doc_id)
    # Send first + middle + last chunks as samples
    sample_ids = [chunks[0]["chunk_id"]] if chunks else []
    if len(chunks) > 2:
        sample_ids.append(chunks[len(chunks)//2]["chunk_id"])
    if len(chunks) > 1:
        sample_ids.append(chunks[-1]["chunk_id"])

    samples = ""
    for cid in sample_ids:
        ch = literary.get_chunk(cid)
        if ch:
            samples += ch["text"][:2000] + "\n\n"

    prompt = f"""Analyze this text and extract project metadata. Respond ONLY with valid JSON.

TEXT SAMPLES ({info['total_words']} words total):
{samples[:6000]}

Return this exact JSON structure (fill every field based on the text):
{{
  "genre": "fiction|non-fiction|screenplay|poetry|academic|technical",
  "brief": "2-3 sentence summary of what this text is about",
  "detected_style": "describe the writing style in 1-2 sentences",
  "detected_intent": "what is the author trying to achieve",
  "tone": "e.g. dark, humorous, formal, intimate, detached",
  "pov": "e.g. first person, third person limited, omniscient",
  "setting": "where and when the story takes place",
  "audience": "who is this written for",
  "themes": ["theme1", "theme2", "theme3"],
  "character_arcs": [
    {{"name": "Character Name", "arc": "brief description of their journey/role"}},
  ],
  "steps": ["major plot point or section 1", "major plot point 2", "..."],
  "target_words": estimated_final_word_count_as_integer
}}"""

    # Use a fast model for this
    for p in _providers:
        if p.available and p.provider in ("groq", "mistral"):
            result = await execute(p, prompt, max_tokens=1500,
                                   system="You are a literary analyst. Respond ONLY with valid JSON, no markdown.")
            if not result.degraded:
                try:
                    m = re.search(r'\{.*\}', result.content, re.DOTALL)
                    if m:
                        data = json.loads(m.group())
                        # Update project with extracted data
                        literary.update_project(project_id, **data)
                        return {"ok": True, "extracted": data}
                except (json.JSONDecodeError, AttributeError):
                    pass
    return {"error": "Could not auto-populate — LLM failed to return valid JSON"}


@app.get("/api/literary/workflow/{project_id}")
async def api_workflow_state(project_id: str):
    return wf.get_workflow_state(project_id)


@app.post("/api/literary/workflow/{project_id}/next")
async def api_workflow_next(
    project_id: str,
    doc_id: str = Form(""),
    mode: str = Form("horizontal"),  # horizontal or vertical
    budget_pct: int = Form(10),      # % of daily credits to spend
):
    """Execute the next logical step in the writing workflow."""
    state = wf.get_workflow_state(project_id)
    if state.get("error"):
        return state

    proj = literary.get_project(project_id)
    doc_text = literary.get_full_text(doc_id) if doc_id else ""

    # Build prompt based on mode
    if mode == "vertical":
        prompt = wf.build_vertical_prompt(proj, state, doc_text, budget_pct)
    else:
        prompt = wf.build_horizontal_prompt(proj, state, doc_text, budget_pct)

    # Inject craft techniques
    craft = build_craft_prompt(
        proj.get("genre", "fiction"),
        ["structure"] if mode == "vertical" else ["prose"],
        proj.get("style", "") or proj.get("detected_style", "")
    )
    if craft:
        prompt += "\n\n" + craft

    # Pick model based on budget — cheap for low budget, best for high
    provider = None
    if budget_pct <= 5:
        for p in _providers:
            if p.available and p.provider in ("groq", "cloudflare"):
                provider = p; break
    elif budget_pct <= 25:
        for p in _providers:
            if p.available and p.provider in ("groq", "mistral", "deepseek"):
                provider = p; break

    if not provider:
        for p in _providers:
            if p.available:
                provider = p; break

    if not provider:
        return {"error": "No providers available"}

    max_tokens = min(8192, max(1000, int(budget_pct * 80)))
    result = await execute(provider, prompt, max_tokens=max_tokens,
                           system="You are a master novelist/writer. Write with precision, depth, and craft.")

    response = {
        "state": state,
        "mode": mode,
        "budget_pct": budget_pct,
        "model": result.provider_name,
        "content": result.content if not result.degraded else f"Failed: {result.content}",
        "word_count": len(result.content.split()) if not result.degraded else 0,
        "cost": result.cost,
    }

    # If outline phase and successful, try to auto-update project
    if not result.degraded and state["phase"] == "outline":
        # Try to extract structure from the response
        try:
            fd = {"steps": [l.strip().lstrip("0123456789.-) ") for l in result.content.split("\n")
                            if l.strip() and any(l.strip().startswith(str(i)) for i in range(1, 20))]}
            if fd["steps"]:
                literary.update_project(project_id, steps=fd["steps"][:15])
        except Exception:
            pass

    # Mark step complete if we were expanding a specific step
    if not result.degraded and state["phase"] == "expand" and state["current_step_idx"] < state["steps_total"]:
        wf.mark_step_complete(project_id, state["current_step_idx"])

    # Compute next state
    response["next_state"] = wf.get_workflow_state(project_id)
    return response


@app.post("/api/literary/{doc_id}/iterate")
async def api_literary_iterate(
    doc_id: str,
    project_id: str = Form(""),
    resume: bool = Form(False),
):
    """Run iterative analysis: chunk by chunk, saving progress. Resumable."""
    info = literary.get_document_info(doc_id)
    if not info:
        return {"error": "Document not found"}

    proj = literary.get_project(project_id) if project_id else None
    state = (proj or {}).get("iteration_state", {}) if resume else {}
    completed_chunks = state.get("completed", [])
    results_so_far = state.get("results", [])

    chunks = literary.get_chunks(doc_id)
    remaining = [c for c in chunks if c["chunk_id"] not in completed_chunks]

    if not remaining:
        return {"status": "complete", "message": "All chunks processed", "results": results_so_far}

    # Process next chunk
    chunk = remaining[0]
    chunk_data = literary.get_chunk(chunk["chunk_id"])
    if not chunk_data:
        return {"error": f"Chunk {chunk['chunk_id']} not found"}

    project_brief = literary.get_project_brief_prompt(project_id) if project_id else ""
    craft_ref = ""
    if proj:
        from intello.craft import build_craft_prompt
        craft_ref = build_craft_prompt(proj.get("genre", "fiction"), ["prose"], proj.get("style", ""))

    chunk_prompt = f"""{f"PROJECT:{chr(10)}{project_brief}{chr(10)}{chr(10)}" if project_brief else ""}{craft_ref}

Analyze this section (chunk {len(completed_chunks)+1}/{len(chunks)}):
Chapter: {chunk_data['chapter']}
Lines {chunk_data['start_line']}-{chunk_data['end_line']}

TEXT:
{chunk_data['text']}

Provide:
1. Quality assessment (1-2 sentences)
2. Specific issues found
3. Concrete edits in format: EDIT LINE X-Y: [text] — REASON: [why]
4. How this section serves the overall work

Be surgical and specific."""

    result = await execute(
        next((p for p in _providers if p.available and p.provider in ("groq", "cloudflare", "mistral")), _providers[0]),
        chunk_prompt, max_tokens=2000,
        system="You are a literary editor doing a line-by-line review. Be precise."
    )

    chunk_result = {
        "chunk_id": chunk["chunk_id"],
        "chapter": chunk["chapter"],
        "lines": f"{chunk['start_line']}-{chunk['end_line']}",
        "analysis": result.content if not result.degraded else f"Failed: {result.content}",
        "model": result.model_id,
    }

    # Parse edits
    if not result.degraded:
        for m in re.finditer(r'EDIT LINE[S]? (\d+)-(\d+):\s*(.+?)(?:\s*—\s*REASON:\s*(.+?))?(?:\n|$)', result.content):
            literary.propose_edit(doc_id, "replace", int(m.group(1)), int(m.group(2)),
                                  m.group(3).strip(), m.group(4) or "", result.model_id)

    completed_chunks.append(chunk["chunk_id"])
    results_so_far.append(chunk_result)

    # Save progress
    new_state = {"completed": completed_chunks, "results": results_so_far,
                 "total_chunks": len(chunks), "last_updated": time.time()}
    if project_id:
        literary.update_project(project_id, iteration_state=new_state)

    return {
        "status": "in_progress",
        "progress": f"{len(completed_chunks)}/{len(chunks)}",
        "chunk_result": chunk_result,
        "remaining": len(remaining) - 1,
    }


@app.post("/api/literary/ingest")
async def api_literary_ingest(
    title: str = Form(""),
    file: Optional[UploadFile] = File(None),
    text: Optional[str] = Form(None),
    gdrive_url: Optional[str] = Form(None),
    project_id: Optional[str] = Form(""),
):
    """Ingest a document for literary analysis. Supports .txt, .md, .pdf, .epub."""
    import tempfile
    fname = title
    doc_id = None

    if file and file.filename:
        fname = fname or file.filename
        ext = file.filename.rsplit(".", 1)[-1].lower() if "." in file.filename else ""
        doc_id = fname.replace(" ", "_").lower()[:40] + f"_{int(time.time())}"

        if ext == "pdf":
            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
                f.write(await file.read())
                tmp = f.name
            result = literary.ingest_pdf(doc_id, tmp, fname, project_id or "")
            os.unlink(tmp)
            return result
        elif ext == "epub":
            with tempfile.NamedTemporaryFile(suffix=".epub", delete=False) as f:
                f.write(await file.read())
                tmp = f.name
            result = literary.ingest_epub(doc_id, tmp, fname, project_id or "")
            os.unlink(tmp)
            return result
        else:
            content = (await file.read()).decode("utf-8", errors="replace")
    elif text:
        content = text
        fname = fname or "pasted_text"
    elif gdrive_url and gdrive_url.strip():
        if gdrive.is_authenticated():
            content = gdrive.fetch_private(gdrive_url.strip())
        else:
            content = await gdrive.fetch_public(gdrive_url.strip())
        fname = fname or "gdrive_doc"
    else:
        return {"error": "No content provided"}

    if not content or len(content) < 50:
        return {"error": "Content too short (min 50 chars)"}

    if not doc_id:
        doc_id = fname.replace(" ", "_").lower()[:40] + f"_{int(time.time())}"
    return literary.ingest_document(doc_id, content, fname, project_id or "")


@app.get("/api/literary/documents")
async def api_literary_documents():
    with literary._db() as conn:
        rows = conn.execute("SELECT doc_id, title, total_lines, total_words, created_at FROM documents ORDER BY created_at DESC").fetchall()
    return [dict(r) for r in rows]


@app.get("/api/literary/{doc_id}")
async def api_literary_doc(doc_id: str):
    info = literary.get_document_info(doc_id)
    if not info:
        return {"error": "Document not found"}
    return {
        "info": info,
        "structure": literary.get_structure(doc_id),
        "chunks": literary.get_chunks(doc_id),
        "pacing": literary.get_pacing_data(doc_id, window=max(5, info["total_lines"] // 20)),
        "characters": literary.get_characters(doc_id),
        "threads": literary.get_threads(doc_id),
    }


@app.get("/api/literary/{doc_id}/lines")
async def api_literary_lines(doc_id: str, start: int = 1, end: int = 50):
    return literary.get_lines(doc_id, start, end)


@app.get("/api/literary/{doc_id}/edits")
async def api_literary_edits(doc_id: str):
    return literary.get_pending_edits(doc_id)


# --- Writing Tools (Sudowrite-style) ---

@app.post("/api/tools/transform")
async def api_tool_transform(
    text: str = Form(...),
    tool: str = Form(...),  # show_not_tell, describe, tone_shift, shrink, first_draft, brainstorm
    context: str = Form(""),
    target: str = Form(""),  # tone name, shrink format, brainstorm category
    genre: str = Form("fiction"),
    word_count: int = Form(1000),
    style: str = Form(""),
):
    """Universal writing tool endpoint."""
    prompts = {
        "show_not_tell": wt.show_not_tell(text, context),
        "describe": wt.describe_senses(text, context),
        "tone_shift": wt.tone_shift(text, target or "darker"),
        "shrink": wt.shrink_ray(text, target or "blurb"),
        "first_draft": wt.first_draft(text, style, word_count),
        "brainstorm": wt.brainstorm(text, target or "plot", genre),
    }
    prompt = prompts.get(tool)
    if not prompt:
        return {"error": f"Unknown tool: {tool}"}

    # Pick a good model
    provider = None
    for p in _providers:
        if p.available and p.provider in ("groq", "mistral", "google"):
            provider = p
            break
    if not provider:
        for p in _providers:
            if p.available:
                provider = p
                break
    if not provider:
        return {"error": "No providers available"}

    result = await execute(provider, prompt, max_tokens=4096,
                           system="You are a master fiction editor and writing coach. Be specific and craft-aware.")
    return {
        "tool": tool,
        "result": result.content if not result.degraded else f"Failed: {result.content}",
        "model": result.provider_name,
        "cost": result.cost,
        "word_count": len(result.content.split()) if not result.degraded else 0,
    }


@app.post("/api/tools/beta-read")
async def api_tool_beta_read(text: str = Form(...)):
    """Run 3 AI beta readers in parallel with different perspectives."""
    import asyncio
    reader_types = ["casual", "craft", "market"]

    # Pick 3 different models for diversity
    models = []
    seen = set()
    for p in _providers:
        if p.available and p.provider not in seen and ratelimit.remaining(p.model_id, p.daily_limit) != 0:
            models.append(p)
            seen.add(p.provider)
            if len(models) >= 3:
                break

    if not models:
        return {"error": "Not enough providers available"}

    tasks = []
    for i, rtype in enumerate(reader_types):
        model = models[i % len(models)]
        prompt = wt.beta_reader_prompt(text, rtype)
        tasks.append(execute(model, prompt, max_tokens=2000,
                             system=f"You are a {rtype} reader. Give honest, detailed feedback."))

    results = await asyncio.gather(*tasks)
    readers = []
    for rtype, result in zip(reader_types, results):
        readers.append({
            "type": rtype,
            "model": result.provider_name,
            "feedback": result.content if not result.degraded else "Failed",
            "cost": result.cost,
        })

    return {"readers": readers, "total_cost": sum(r["cost"] for r in readers)}


@app.post("/api/literary/{doc_id}/analyze")
async def api_literary_analyze(doc_id: str, focus: str = Form("full")):
    """Run multi-model literary analysis on a document."""
    info = literary.get_document_info(doc_id)
    if not info:
        return {"error": "Document not found"}

    structure = literary.get_structure(doc_id)
    pacing = literary.get_pacing_data(doc_id, window=30)
    chunks = literary.get_chunks(doc_id)

    # Build analysis prompt with structure + pacing summary
    struct_summary = "\n".join(
        f"  {json.loads(s['metadata']).get('title','?')} (lines {s['start_line']}-{s['end_line']})"
        for s in structure
    )
    pacing_summary = ""
    for p in pacing:
        tension_label = "🔴 HIGH" if p["tension"] > 1.5 else "🟡 MED" if p["tension"] > 0.5 else "🟢 LOW"
        pacing_summary += f"  Lines {p['start_line']}-{p['end_line']}: tension={tension_label}, dialogue={p['dialogue_ratio']:.0%}, avg_sent={p['avg_sentence_len']:.1f} words\n"

    # Pick 2-3 chunks to send as samples (beginning, middle, end)
    sample_ids = []
    if chunks:
        sample_ids = [chunks[0]["chunk_id"]]
        if len(chunks) > 2:
            sample_ids.append(chunks[len(chunks) // 2]["chunk_id"])
        if len(chunks) > 1:
            sample_ids.append(chunks[-1]["chunk_id"])

    samples = ""
    for cid in sample_ids:
        ch = literary.get_chunk(cid)
        if ch:
            samples += f"\n--- {ch['chapter']} (lines {ch['start_line']}-{ch['end_line']}) ---\n{ch['text'][:3000]}\n"

    # Get project brief if linked
    project_brief = ""
    project_genre = "fiction"
    project_style = ""
    project_target = 0
    doc_project_id = info.get("project_id") or ""
    if isinstance(doc_project_id, str) and doc_project_id:
        project_brief = literary.get_project_brief_prompt(doc_project_id)
        proj = literary.get_project(doc_project_id)
        if proj:
            project_genre = proj.get("genre", "fiction")
            project_style = proj.get("style", "")
            project_target = proj.get("target_words", 0)

    # Detect issues from pacing data for craft reference
    pacing_issues = []
    for p in pacing:
        if p["tension"] < 0.3:
            pacing_issues.append("slow pacing")
        if p["tension"] > 2:
            pacing_issues.append("fast pacing")
        if p["dialogue_ratio"] < 0.05:
            pacing_issues.append("flat — no dialogue")
    pacing_issues = list(set(pacing_issues))[:5]

    craft_ref = build_craft_prompt(project_genre, pacing_issues, project_style)

    # Word count check
    wc_note = ""
    if project_target:
        wc = check_word_count(literary.get_full_text(doc_id), project_target)
        wc_note = f"\nWORD COUNT: {wc['actual']} / {wc['target']} target ({wc['verdict']})"

    analysis_prompt = f"""Analyze this document as a literary expert.

{f"PROJECT CONTEXT:{chr(10)}{project_brief}{chr(10)}" if project_brief else ""}{wc_note}

DOCUMENT: "{info['title']}" — {info['total_lines']} lines, {info['total_words']} words

STRUCTURE:
{struct_summary}

PACING ANALYSIS:
{pacing_summary}

{craft_ref}

SAMPLE EXCERPTS:
{samples}

Provide a detailed literary analysis covering:
1. STRUCTURE: Is the chapter/scene organization effective? Where should breaks be added/removed?
2. PACING: Where is the narrative too slow or too fast? Which sections need tightening or expansion?
3. PROSE QUALITY: Comment on sentence variety, word choice, rhythm. Quote specific lines.
4. WORD COUNT: Is the current length appropriate for the target? Which sections are bloated or too thin?
5. SUGGESTIONS: For each issue, provide SPECIFIC edits in this format:
   - EDIT LINE X-Y: [replacement text] — REASON: [why]
   - INSERT AFTER LINE X: [new text] — REASON: [why]

IMPORTANT: When you suggest new or replacement text, it MUST hit the word count implied by the edit.
Do NOT write "a 500-word paragraph about X" — actually WRITE the 500 words.

Be brutally honest. This writer wants to produce super literature, not hear compliments."""

    # Run in deep mode for best quality
    from intello.pipeline import run_deep
    pipe = await run_deep(analysis_prompt, _providers)

    # Word count verification of the analysis output
    wc_result = None
    if project_target and pipe.final and not pipe.final.degraded:
        wc_result = check_word_count(literary.get_full_text(doc_id), project_target)

    # Parse any edit suggestions from the response
    edits_proposed = 0
    if pipe.final and not pipe.final.degraded:
        # Try to extract EDIT LINE patterns
        for m in re.finditer(r'EDIT LINE[S]? (\d+)-(\d+):\s*(.+?)(?:\s*—\s*REASON:\s*(.+?))?(?:\n|$)',
                             pipe.final.content):
            literary.propose_edit(doc_id, "replace", int(m.group(1)), int(m.group(2)),
                                  m.group(3).strip(), m.group(4) or "", pipe.final.model_id)
            edits_proposed += 1
        for m in re.finditer(r'INSERT AFTER LINE (\d+):\s*(.+?)(?:\s*—\s*REASON:\s*(.+?))?(?:\n|$)',
                             pipe.final.content):
            literary.propose_edit(doc_id, "insert", int(m.group(1)), int(m.group(1)),
                                  m.group(2).strip(), m.group(3) or "", pipe.final.model_id)
            edits_proposed += 1

    return {
        "analysis": pipe.final.content if pipe.final else "Analysis failed",
        "edits_proposed": edits_proposed,
        "pipeline_steps": pipe.steps_log,
        "cost": pipe.total_cost,
        "word_count": wc_result,
        "craft_techniques_used": len(pacing_issues),
    }


@app.post("/api/literary/{doc_id}/edit/{edit_id}/apply")
async def api_literary_apply_edit(doc_id: str, edit_id: int):
    ok = literary.apply_edit(edit_id)
    return {"ok": ok}


@app.post("/api/literary/{doc_id}/edit/{edit_id}/reject")
async def api_literary_reject_edit(doc_id: str, edit_id: int):
    literary.reject_edit(edit_id)
    return {"ok": True}


@app.post("/api/literary/{doc_id}/append")
async def api_literary_append(doc_id: str, text: str = Form(...)):
    """Append text to a document (from workflow output, etc.)."""
    info = literary.get_document_info(doc_id)
    if not info:
        return {"error": "Document not found"}
    current = literary.get_full_text(doc_id)
    new_text = current + "\n\n" + text
    # Re-ingest with appended content
    result = literary.ingest_document(doc_id, new_text, info["title"],
                                      info.get("project_id", ""))
    return result


@app.get("/api/literary/{doc_id}/export/docx")
async def api_literary_export_docx(doc_id: str):
    """Export document as DOCX."""
    info = literary.get_document_info(doc_id)
    if not info:
        return Response("Not found", status_code=404)

    from docx import Document as DocxDocument
    from docx.shared import Pt
    import tempfile

    doc = DocxDocument()
    doc.add_heading(info["title"], 0)

    with literary._db() as conn:
        lines = conn.execute("SELECT line_num, text, chapter FROM lines WHERE doc_id=? ORDER BY line_num",
                             (doc_id,)).fetchall()

    current_chapter = ""
    for line in lines:
        if line["chapter"] != current_chapter:
            current_chapter = line["chapter"]
            doc.add_heading(current_chapter, level=1)
        elif line["text"].strip():
            doc.add_paragraph(line["text"])

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        doc.save(f.name)
        tmp = f.name

    with open(tmp, "rb") as f:
        content = f.read()
    os.unlink(tmp)

    return Response(content, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    headers={"Content-Disposition": f'attachment; filename="{info["title"]}.docx"'})


@app.get("/api/literary/{doc_id}/export", response_class=HTMLResponse)
async def api_literary_export(doc_id: str):
    """Generate a rich, editable HTML report of the full literary analysis."""
    info = literary.get_document_info(doc_id)
    if not info:
        return HTMLResponse("<h1>Document not found</h1>", status_code=404)

    structure = literary.get_structure(doc_id)
    characters = literary.get_characters(doc_id)
    pacing = literary.get_pacing_data(doc_id, window=max(5, info["total_lines"] // 25))
    threads = literary.get_threads(doc_id)
    edits = literary.get_pending_edits(doc_id)

    with literary._db() as conn:
        lines = conn.execute("SELECT line_num, text, chapter FROM lines WHERE doc_id=? ORDER BY line_num",
                             (doc_id,)).fetchall()
        pacing_raw = conn.execute("SELECT line_num, tension, dialogue FROM pacing WHERE doc_id=? ORDER BY line_num",
                                  (doc_id,)).fetchall()

    # Build edit lookup by line
    edit_map = {}
    for e in edits:
        for ln in range(e["start_line"], e["end_line"] + 1):
            edit_map[ln] = e

    # Pacing lookup
    pacing_map = {r["line_num"]: dict(r) for r in pacing_raw}

    # Thread color map
    thread_colors = {
        "identity": "#ef4444", "motive": "#f97316", "event": "#eab308", "method": "#22c55e",
        "location": "#06b6d4", "timing": "#3b82f6", "suspense": "#8b5cf6", "secret": "#ec4899",
        "mystery": "#a855f7", "concealment": "#6366f1", "deception": "#f43f5e", "curiosity": "#14b8a6",
        "anomaly": "#f59e0b", "decision": "#10b981", "unknown": "#6b7280", "uncertainty": "#9ca3af",
        "threat": "#dc2626", "promise": "#2563eb", "investigation": "#7c3aed",
    }

    # --- Build HTML ---
    import html as html_mod

    def esc(s):
        return html_mod.escape(str(s))

    # Pacing SVG
    max_t = max((p["tension"] for p in pacing), default=1) or 1
    svg_w, svg_h = 700, 80
    points = []
    for i, p in enumerate(pacing):
        x = (i / max(len(pacing) - 1, 1)) * svg_w
        y = svg_h - (p["tension"] / max_t) * (svg_h - 10)
        points.append(f"{x:.0f},{y:.0f}")
    pacing_svg = (
        f'<svg width="{svg_w}" height="{svg_h}" style="width:100%;height:{svg_h}px">'
        f'<polyline points="{" ".join(points)}" fill="none" stroke="#6366f1" stroke-width="2"/>'
    )
    # Add colored dots
    for i, p in enumerate(pacing):
        x = (i / max(len(pacing) - 1, 1)) * svg_w
        y = svg_h - (p["tension"] / max_t) * (svg_h - 10)
        color = "#ef4444" if p["tension"] > max_t * 0.7 else "#eab308" if p["tension"] > max_t * 0.3 else "#22c55e"
        pacing_svg += f'<circle cx="{x:.0f}" cy="{y:.0f}" r="4" fill="{color}"><title>L{p["start_line"]}-{p["end_line"]} tension={p["tension"]:.1f}</title></circle>'
    pacing_svg += '</svg>'

    # Thread bars SVG
    total = info["total_lines"]
    thread_svg = f'<svg width="100%" height="{len(threads) * 14 + 4}" style="width:100%">'
    for i, t in enumerate(threads):
        x1 = (t["start_line"] / total) * 100
        w = max(1, ((t["end_line"] - t["start_line"]) / total) * 100)
        color = thread_colors.get(t["category"], "#6b7280")
        opacity = "0.5" if t["resolved"] else "0.9"
        status = "✅" if t["resolved"] else "❓"
        thread_svg += (f'<rect x="{x1:.1f}%" y="{i * 14}" width="{w:.1f}%" height="10" rx="3" '
                       f'fill="{color}" opacity="{opacity}">'
                       f'<title>{status} {esc(t["category"])}: {esc(t["description"][:80])}</title></rect>')
    thread_svg += '</svg>'

    # Annotated text
    text_html = ""
    current_chapter = ""
    for row in lines:
        ln = row["line_num"]
        txt = esc(row["text"])
        p = pacing_map.get(ln, {})
        tension = p.get("tension", 0) if p else 0
        is_dialogue = p.get("dialogue", 0) if p else 0

        # Chapter header
        if row["chapter"] != current_chapter:
            current_chapter = row["chapter"]
            text_html += f'<h3 style="color:#6366f1;margin:24px 0 8px;page-break-before:auto" id="line-{ln}">{txt}</h3>\n'
            continue

        # Line styling
        style = ""
        cls = ""
        if tension > 1.5:
            style = "border-left:3px solid #ef4444;padding-left:8px;"
            cls = "high-tension"
        elif is_dialogue:
            style = "color:#3b82f6;"

        # Edit annotation
        edit_note = ""
        if ln in edit_map:
            e = edit_map[ln]
            edit_note = (f'<span style="background:#22c55e22;border:1px solid #22c55e;border-radius:4px;'
                         f'padding:2px 6px;font-size:.8rem;margin-left:8px" contenteditable="false">'
                         f'✏️ {esc(e["reason"][:60])}</span>')

        text_html += (f'<div style="display:flex;gap:12px;{style}" id="line-{ln}">'
                      f'<span style="color:#71717a;font-size:.75rem;min-width:35px;text-align:right;'
                      f'font-family:monospace;user-select:none">{ln}</span>'
                      f'<span contenteditable="true" style="flex:1">{txt}</span>'
                      f'{edit_note}</div>\n')

    # Character summary
    char_html = "".join(
        f'<span style="display:inline-block;background:#1a1d27;border:1px solid #2a2d3a;'
        f'padding:3px 10px;border-radius:12px;margin:3px;font-size:.85rem">'
        f'{esc(c["name"])} <span style="color:#71717a;font-size:.75rem">{c["mentions"]}×</span></span>'
        for c in characters
    )

    # Thread descriptions
    thread_desc_html = ""
    for t in threads:
        color = thread_colors.get(t["category"], "#6b7280")
        status = "✅ Resolved" if t["resolved"] else "❓ Open"
        thread_desc_html += (
            f'<div style="border-left:4px solid {color};padding:6px 12px;margin:4px 0;'
            f'background:#0f1117;border-radius:0 6px 6px 0;font-size:.85rem">'
            f'<span style="color:{color};font-size:.7rem;text-transform:uppercase;font-weight:600">{esc(t["category"])}</span> '
            f'<span style="color:#71717a;font-size:.75rem">L{t["start_line"]}–{t["end_line"]}</span> '
            f'<span style="font-size:.75rem">{status}</span><br>'
            f'{esc(t["description"])}</div>'
        )

    # Assemble
    report = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<title>{esc(info['title'])} — Literary Analysis</title>
<style>
  @media print {{ @page {{ margin: 2cm; }} .no-print {{ display: none; }} }}
  body {{ font-family: Georgia, serif; max-width: 900px; margin: 0 auto; padding: 24px;
          background: #fff; color: #1a1a1a; line-height: 1.8; }}
  h1 {{ font-size: 1.8rem; border-bottom: 2px solid #6366f1; padding-bottom: 8px; }}
  h2 {{ font-size: 1.2rem; color: #6366f1; margin: 32px 0 12px; }}
  .stats {{ color: #666; font-size: .9rem; margin-bottom: 24px; }}
  .section {{ margin: 24px 0; padding: 16px; background: #f8f9fa; border-radius: 8px; }}
  [contenteditable=true]:focus {{ outline: 2px solid #6366f1; border-radius: 4px; }}
  [contenteditable=true]:hover {{ background: #f0f0ff; }}
  .toolbar {{ position: sticky; top: 0; background: #fff; padding: 8px 0; border-bottom: 1px solid #ddd;
              z-index: 10; display: flex; gap: 8px; }}
  .toolbar button {{ padding: 6px 14px; border: 1px solid #ddd; border-radius: 6px; cursor: pointer;
                     background: #fff; font-size: .85rem; }}
  .toolbar button:hover {{ background: #f0f0ff; }}
</style>
</head>
<body>

<div class="toolbar no-print">
  <button onclick="window.print()">🖨️ Print / PDF</button>
  <button onclick="downloadHTML()">💾 Save HTML</button>
  <span style="color:#999;font-size:.85rem;margin-left:auto">Click any text to edit directly</span>
</div>

<h1>{esc(info['title'])}</h1>
<div class="stats">{info['total_lines']} lines · {info['total_words']} words · {len(structure)} chapters · {len(characters)} characters · {len(threads)} narrative threads</div>

<h2>📖 Structure</h2>
<div class="section">
{"".join(f'<div><a href="#line-{s["start_line"]}" style="color:#6366f1">{esc(json.loads(s["metadata"])["title"])}</a> <span style="color:#999;font-size:.85rem">lines {s["start_line"]}–{s["end_line"]}</span></div>' for s in structure)}
</div>

<h2>👤 Characters</h2>
<div class="section">{char_html or '<span style="color:#999">No characters detected</span>'}</div>

<h2>📊 Pacing Curve</h2>
<div class="section">
  <div style="display:flex;justify-content:space-between;font-size:.75rem;color:#999"><span>Start</span><span>🟢 Low — 🟡 Medium — 🔴 High tension</span><span>End</span></div>
  {pacing_svg}
</div>

<h2>🧵 Narrative Threads</h2>
<div class="section">
  {thread_svg}
  <div style="margin-top:12px">{thread_desc_html}</div>
</div>

<h2>📝 Annotated Text</h2>
<div style="font-size:.95rem;line-height:1.9">
{text_html}
</div>

<script>
function downloadHTML() {{
  const html = document.documentElement.outerHTML;
  const blob = new Blob([html], {{type: 'text/html'}});
  const a = document.createElement('a');
  a.href = URL.createObjectURL(blob);
  a.download = '{esc(info["title"]).replace("'", "")}_analysis.html';
  a.click();
}}
</script>
</body>
</html>"""

    return HTMLResponse(report)


# --- Context Compression ---

async def _compress_context(conv_id: str):
    """Summarize old messages using a cheap fast model."""
    msgs = memory.get_messages(conv_id, limit=50)
    if len(msgs) < 10:
        return
    old_summary = memory.get_summary(conv_id)
    text_parts = []
    if old_summary:
        text_parts.append(f"Previous summary: {old_summary}")
    for m in msgs[:-5]:  # Keep last 5 raw, compress the rest
        text_parts.append(f"{m['role']}: {m['content'][:500]}")

    compress_prompt = (
        "Compress this conversation into a concise summary (max 300 words). "
        "Capture key topics, decisions, user preferences, and any important context. "
        "Be factual and dense.\n\n" + "\n".join(text_parts)
    )

    # Pick cheapest available model for compression
    for p in _providers:
        if p.available and p.provider in ("groq", "cloudflare", "mistral"):
            result = await execute(p, compress_prompt, max_tokens=500,
                                   system="You are a conversation summarizer. Be concise and factual.")
            if not result.degraded:
                memory.set_summary(conv_id, result.content)
                return


# --- Prompt handling ---

@app.post("/api/prompt")
async def api_prompt(
    prompt: str = Form(...),
    gdrive_url: Optional[str] = Form(None),
    file: Optional[UploadFile] = File(None),
    confirm_paid: bool = Form(False),
    mode: Optional[str] = Form(None),
    conversation_id: Optional[str] = Form(None),
    request: Request = None,
):
    # User-based model filtering
    user = _get_user(request) if request else "anonymous"
    user_providers = filter_providers_for_user(_providers, user)

    # Load user prefs
    prefs = memory.get_prefs()
    effective_mode = mode or prefs.get("default_mode", "auto")

    # System prompt from prefs
    sys_prompt = prefs.get("custom_system_prompt") or None
    if not sys_prompt:
        tone = prefs.get("tone", "brutally_honest")
        if tone != "brutally_honest":
            sys_prompt = f"You are a helpful AI assistant. Respond in a {tone} tone."

    # Conversation setup
    conv_id = conversation_id or str(uuid.uuid4())
    memory.create_conversation(conv_id)

    # Build prompt with file context
    parts = []
    if file and file.filename:
        content = (await file.read()).decode("utf-8", errors="replace")
        parts.append(f"--- Attached file: {file.filename} ---\n{content[:50_000]}")
    if gdrive_url and gdrive_url.strip():
        url = gdrive_url.strip()
        if gdrive.is_authenticated():
            text = gdrive.fetch_private(url)
        else:
            text = await gdrive.fetch_public(url)
        parts.append(f"--- Google Drive file ---\n{text}")

    # Add conversation context
    context = memory.build_context(conv_id, max_messages=10)
    if context:
        parts.append(f"--- Conversation context ---\n{context}")

    parts.append(prompt)
    full_prompt = "\n\n".join(parts)

    # Save user message
    memory.add_message(conv_id, "user", prompt)

    # Compress if needed (async, don't block)
    if memory.needs_compression(conv_id):
        asyncio.create_task(_compress_context(conv_id))

    plan = build_plan(full_prompt, user_providers)
    plan_info = {
        "task_type": plan.task_type.value,
        "estimated_tokens": plan.estimated_tokens,
        "primary": plan.primary.name if plan.primary else None,
        "fallbacks": [f.name for f in plan.fallbacks],
        "degraded": plan.degraded,
        "missing_keys": plan.missing_keys,
        "estimated_cost": plan.estimated_cost,
        "reasoning": plan.reasoning,
        "is_paid": plan.primary.tier == Tier.PAID if plan.primary else False,
        "conversation_id": conv_id,
    }

    if plan.degraded or not plan.primary:
        return {"plan": plan_info, "needs_confirmation": False,
                "error": "No providers available. Supply API keys to continue."}

    # Check cache (skip for deep mode — user wants fresh multi-model analysis)
    if effective_mode != "deep":
        cached = cache.get_cached(prompt, plan.task_type.value)
        if cached:
            memory.add_message(conv_id, "assistant", cached["response"], cached["model"], 0)
            plan_info["reasoning"] += "\n⚡ Cache hit — reusing previous response"
            return {
                "plan": plan_info, "needs_confirmation": False, "mode": "cached",
                "response": {
                    "provider": cached["provider"], "model": cached["model"],
                    "content": cached["response"],
                    "input_tokens": 0, "output_tokens": 0, "cost": 0,
                },
            }

    # Decide mode
    use_deep = effective_mode == "deep" or (
        effective_mode == "auto" and (plan.estimated_tokens > 500 or plan.task_type.value in ("analysis", "code", "long_context"))
    )

    # --- Debate mode ---
    if effective_mode == "debate":
        plan_info["reasoning"] += "\n⚔️ Debate mode: positions → challenges → verdict"
        debate = await run_debate(full_prompt, user_providers, system=sys_prompt)
        if debate.verdict.get("content"):
            memory.add_message(conv_id, "assistant", debate.verdict["content"],
                               debate.verdict.get("model_id", ""), debate.total_cost)
        return {
            "plan": plan_info, "needs_confirmation": False, "mode": "debate",
            "debate": {
                "log": debate.log,
                "positions": debate.positions,
                "challenges": debate.challenges,
                "verdict": debate.verdict,
            },
            "response": {
                "provider": debate.verdict.get("judge", "debate"),
                "model": debate.verdict.get("model_id", ""),
                "content": debate.verdict.get("content", "Debate failed"),
                "input_tokens": 0, "output_tokens": 0,
                "cost": debate.total_cost,
            } if debate.verdict.get("content") else None,
            "error": "Debate failed" if not debate.verdict.get("content") else None,
        }

    if use_deep:
        plan_info["reasoning"] += "\n🔬 Deep mode: multi-LLM draft → cross-review → synthesis"
        t0 = time.time()
        pipe = await run_deep(full_prompt, user_providers, system=sys_prompt)
        latency = time.time() - t0

        if pipe.final and not pipe.final.degraded:
            memory.add_message(conv_id, "assistant", pipe.final.content, pipe.final.model_id, pipe.total_cost)
            memory.record_model_result(pipe.final.model_id, plan.task_type.value, True, latency)
            cache.store(prompt, plan.task_type.value, pipe.final.content,
                        pipe.final.provider_name, pipe.final.model_id, pipe.total_cost)

        return {
            "plan": plan_info, "needs_confirmation": False, "mode": "deep",
            "pipeline": {
                "steps": pipe.steps_log,
                "drafts": [{"provider": d.provider_name, "model": d.model_id,
                            "content": d.content, "cost": d.cost, "degraded": d.degraded}
                           for d in pipe.draft_responses],
                "reviews": [{"provider": r.provider_name, "content": r.content, "cost": r.cost}
                            for r in pipe.reviews],
            },
            "response": {
                "provider": pipe.final.provider_name if pipe.final else "none",
                "model": pipe.final.model_id if pipe.final else "",
                "content": pipe.final.content if pipe.final else "Pipeline failed",
                "input_tokens": pipe.final.input_tokens if pipe.final else 0,
                "output_tokens": pipe.final.output_tokens if pipe.final else 0,
                "cost": pipe.total_cost,
            } if pipe.final else None,
            "error": "Pipeline failed" if not pipe.final else None,
        }

    # Fast mode
    if plan.primary.tier == Tier.PAID and not confirm_paid:
        return {"plan": plan_info, "needs_confirmation": True}

    # --- Auto-chaining: check if prompt needs decomposition ---
    if effective_mode == "auto" and plan.estimated_tokens > 30:
        complexity = await analyze_complexity(prompt, user_providers)
        if complexity.get("chain") and complexity.get("steps"):
            plan_info["reasoning"] += "\n🔗 Chain mode: decomposed into " + str(len(complexity["steps"])) + " sub-tasks"
            chain_result = await execute_chain(prompt, complexity["steps"], user_providers, system=sys_prompt)
            final = chain_result["final"]
            if final.get("content"):
                memory.add_message(conv_id, "assistant", final["content"], final.get("model", ""), final.get("cost", 0))
                cache.store(prompt, plan.task_type.value, final["content"],
                            final.get("provider", ""), final.get("model", ""), final.get("cost", 0))
            return {
                "plan": plan_info, "needs_confirmation": False, "mode": "chain",
                "chain_steps": chain_result["steps"],
                "response": {
                    "provider": final.get("provider", "chain"),
                    "model": final.get("model", "synthesis"),
                    "content": final.get("content", "Chain failed"),
                    "input_tokens": 0, "output_tokens": 0,
                    "cost": final.get("cost", 0),
                },
            }

    # --- Single-shot with tool support + guardrails ---
    chain = [plan.primary] + plan.fallbacks
    tool_prompt = full_prompt + TOOL_PROMPT_SUFFIX

    for provider in chain:
        if not provider or not provider.available:
            continue
        t0 = time.time()
        result = await execute(provider, tool_prompt, system=sys_prompt)
        latency = time.time() - t0

        if result.degraded:
            memory.record_model_result(provider.model_id, plan.task_type.value, False, latency)
            continue

        # Check for tool calls
        tool_call = detect_tool_call(result.content)
        if tool_call:
            tool_result = await execute_tool(tool_call)
            # Re-query with tool result
            followup = (f"{full_prompt}\n\n"
                        f"Tool '{tool_call['tool']}' returned:\n{tool_result}\n\n"
                        f"Now answer the original question using this information.")
            result = await execute(provider, followup, system=sys_prompt)
            if result.degraded:
                continue

        # Guardrails check
        confidence = check_confidence(result.content)
        guardrail_info = None
        if confidence["needs_reroute"] and len(chain) > 1:
            # Try a different model
            plan_info["reasoning"] += f"\n⚠️ Low confidence ({confidence['confidence']}) — rerouting"
            for alt in chain:
                if alt and alt.available and alt.model_id != provider.model_id:
                    alt_result = await execute(alt, full_prompt, system=sys_prompt)
                    if not alt_result.degraded:
                        alt_conf = check_confidence(alt_result.content)
                        if alt_conf["confidence"] > confidence["confidence"]:
                            result = alt_result
                            confidence = alt_conf
                            break
        if confidence["issues"]:
            guardrail_info = confidence

        memory.add_message(conv_id, "assistant", result.content, result.model_id, result.cost)
        memory.record_model_result(result.model_id, plan.task_type.value, True, latency)
        cache.store(prompt, plan.task_type.value, result.content,
                    result.provider_name, result.model_id, result.cost)
        resp = {
            "plan": plan_info, "needs_confirmation": False, "mode": "fast",
            "response": {
                "provider": result.provider_name, "model": result.model_id,
                "content": result.content, "input_tokens": result.input_tokens,
                "output_tokens": result.output_tokens, "cost": result.cost,
            },
        }
        if guardrail_info:
            resp["guardrails"] = guardrail_info
        if tool_call:
            resp["tool_used"] = {"tool": tool_call["tool"], "args": tool_call.get("args", {})}
        return resp

    return {"plan": plan_info, "needs_confirmation": False, "error": "All providers failed."}


@app.get("/", response_class=HTMLResponse)
async def index():
    with open(os.path.join(os.path.dirname(__file__), "static", "index.html")) as f:
        return f.read()


@app.get("/literary", response_class=HTMLResponse)
async def literary_page():
    with open(os.path.join(os.path.dirname(__file__), "static", "literary.html")) as f:
        return f.read()


@app.get("/corkboard", response_class=HTMLResponse)
async def corkboard_page():
    with open(os.path.join(os.path.dirname(__file__), "static", "corkboard.html")) as f:
        return f.read()


@app.get("/gdrive", response_class=HTMLResponse)
async def gdrive_page():
    with open(os.path.join(os.path.dirname(__file__), "static", "gdrive.html")) as f:
        return f.read()


# --- Scheduled Tasks (#3) ---

@app.get("/api/scheduler/tasks")
async def api_scheduler_list():
    return scheduler.list_tasks()


@app.post("/api/scheduler/tasks")
async def api_scheduler_create(
    name: str = Form(...), prompt: str = Form(...), schedule: str = Form("daily"),
):
    import uuid as _uuid
    tid = _uuid.uuid4().hex[:8]
    return scheduler.create_task(tid, name, prompt, schedule)


@app.delete("/api/scheduler/tasks/{task_id}")
async def api_scheduler_delete(task_id: str):
    scheduler.delete_task(task_id)
    return {"ok": True}


@app.post("/api/scheduler/run")
async def api_scheduler_run_due():
    """Run all due scheduled tasks now."""
    due = scheduler.get_due_tasks()
    results = []
    for task in due:
        plan = build_plan(task["prompt"], _providers)
        if plan.primary:
            result = await execute(plan.primary, task["prompt"], max_tokens=2000)
            scheduler.record_result(task["task_id"], result.content if not result.degraded else "Failed")
            results.append({"task": task["name"], "status": "done" if not result.degraded else "failed"})
        else:
            results.append({"task": task["name"], "status": "no_provider"})
    return {"ran": len(results), "results": results}


# --- Image Generation (#4) ---

@app.post("/api/v1/image/generate")
async def api_image_gen(
    prompt: str = Form(...), style: str = Form(""),
):
    return await imagegen.generate_image(prompt, _providers, style)


# --- Speech Services (Piper TTS + Groq Whisper STT) ---

from intello import speech

@app.post("/api/v1/voice/transcribe")
async def api_voice_transcribe(
    file: UploadFile = File(...),
    language: str = Form(""),
):
    """Speech-to-text via Groq Whisper. Free: 28,800 sec/day (~240 pages)."""
    audio_bytes = await file.read()
    return await speech.transcribe_groq(audio_bytes, file.filename or "audio.wav", language)


@app.post("/api/v1/voice/synthesize")
async def api_voice_synthesize(
    text: str = Form(...),
    language: str = Form("en"),
):
    """Text-to-speech via Piper (local, EN + FR). Returns WAV audio."""
    if not speech.tts_available():
        return {"error": "Piper TTS not installed"}

    audio = speech.synthesize(text, language)
    if not audio:
        return {"error": f"TTS failed for language '{language}'"}

    return Response(audio, media_type="audio/wav",
                    headers={"Content-Disposition": f"attachment; filename=speech_{language}.wav"})


@app.get("/api/v1/voice/voices")
async def api_voice_list():
    """List available TTS voices."""
    return {
        "tts_available": speech.tts_available(),
        "voices": speech.get_available_voices(),
        "stt_provider": "groq (whisper-large-v3-turbo)",
        "stt_daily_limit": "28,800 seconds (~480 minutes)",
    }


# --- Multi-Document Comparison (#6) ---

@app.post("/api/literary/compare")
async def api_literary_compare(
    doc_id_a: str = Form(...), doc_id_b: str = Form(...),
):
    """Compare two documents — structure, pacing, characters, word count."""
    info_a = literary.get_document_info(doc_id_a)
    info_b = literary.get_document_info(doc_id_b)
    if not info_a or not info_b:
        return {"error": "Document not found"}

    chars_a = literary.get_characters(doc_id_a)
    chars_b = literary.get_characters(doc_id_b)
    struct_a = literary.get_structure(doc_id_a)
    struct_b = literary.get_structure(doc_id_b)
    pacing_a = literary.get_pacing_data(doc_id_a, window=max(5, info_a["total_lines"] // 20))
    pacing_b = literary.get_pacing_data(doc_id_b, window=max(5, info_b["total_lines"] // 20))

    char_names_a = {c["name"] for c in chars_a}
    char_names_b = {c["name"] for c in chars_b}

    return {
        "doc_a": {"title": info_a["title"], "words": info_a["total_words"],
                  "chapters": len(struct_a), "characters": len(chars_a)},
        "doc_b": {"title": info_b["title"], "words": info_b["total_words"],
                  "chapters": len(struct_b), "characters": len(chars_b)},
        "word_diff": info_b["total_words"] - info_a["total_words"],
        "chapter_diff": len(struct_b) - len(struct_a),
        "characters_added": list(char_names_b - char_names_a),
        "characters_removed": list(char_names_a - char_names_b),
        "characters_common": list(char_names_a & char_names_b),
        "avg_tension_a": sum(p["tension"] for p in pacing_a) / len(pacing_a) if pacing_a else 0,
        "avg_tension_b": sum(p["tension"] for p in pacing_b) / len(pacing_b) if pacing_b else 0,
    }


# --- Webhooks (#7) ---

@app.get("/api/webhooks")
async def api_webhooks_list():
    return webhooks.list_webhooks()


@app.post("/api/webhooks")
async def api_webhooks_create(
    name: str = Form(...), action: str = Form("chat"), config: str = Form("{}"),
):
    import uuid as _uuid
    hid = _uuid.uuid4().hex[:8]
    return webhooks.create_webhook(hid, name, action, json.loads(config))


@app.post("/api/webhooks/{hook_id}/trigger")
async def api_webhook_trigger(hook_id: str, request: Request):
    """Trigger a webhook — external services call this."""
    hook = webhooks.get_webhook(hook_id)
    if not hook or not hook["enabled"]:
        return {"error": "Webhook not found or disabled"}

    body = await request.json() if request.headers.get("content-type", "").startswith("application/json") else {}
    prompt = body.get("prompt", hook["config"].get("default_prompt", ""))
    if not prompt:
        return {"error": "No prompt in payload or webhook config"}

    plan = build_plan(prompt, _providers)
    if not plan.primary:
        return {"error": "No providers"}

    result = await execute(plan.primary, prompt, max_tokens=body.get("max_tokens", 2000))
    webhooks.log_trigger(hook_id, body, result.content if not result.degraded else "Failed")

    return {"content": result.content, "provider": result.provider_name,
            "model": result.model_id, "cost": result.cost}


@app.delete("/api/webhooks/{hook_id}")
async def api_webhook_delete(hook_id: str):
    webhooks.delete_webhook(hook_id)
    return {"ok": True}


# --- Version Reconstruction (#9) ---

@app.get("/api/reconstruct/projects")
async def api_recon_projects():
    return recon.list_version_projects()


@app.post("/api/reconstruct/projects")
async def api_recon_create(name: str = Form(...)):
    import uuid as _uuid
    pid = name.replace(" ", "_").lower()[:30] + f"_{int(time.time())}"
    return recon.create_version_project(pid, name)


@app.post("/api/reconstruct/{project_id}/ingest")
async def api_recon_ingest(project_id: str, file: UploadFile = File(...)):
    content = (await file.read()).decode("utf-8", errors="replace")
    return recon.ingest_version(project_id, file.filename, content)


@app.get("/api/reconstruct/{project_id}/versions")
async def api_recon_versions(project_id: str):
    return recon.get_project_versions(project_id)


@app.post("/api/reconstruct/{project_id}/rebuild")
async def api_recon_rebuild(project_id: str):
    return recon.reconstruct(project_id)


@app.get("/api/reconstruct/{project_id}/text")
async def api_recon_text(project_id: str):
    text = recon.get_reconstructed_text(project_id)
    return Response(text, media_type="text/plain")


@app.post("/api/reconstruct/{project_id}/smooth")
async def api_recon_smooth(project_id: str):
    """Use LLM to smooth transitions between sections from different versions."""
    text = recon.get_reconstructed_text(project_id)
    if not text:
        return {"error": "No reconstructed text"}

    prompt = f"""This document was reconstructed from multiple versions. Some sections may have inconsistent tone, tense, or style.

Review the transitions between sections and suggest specific edits to make it read as one cohesive document.
For each issue, specify the exact text to change.

DOCUMENT:
{text[:8000]}

List issues and fixes:"""

    plan = build_plan(prompt, _providers)
    if not plan.primary:
        return {"error": "No providers"}
    result = await execute(plan.primary, prompt, max_tokens=4000)
    return {"suggestions": result.content, "provider": result.provider_name}


# --- Backup/Restore (#7) ---

@app.get("/api/backup")
async def api_backup():
    """Download all SQLite databases as a tar archive."""
    import tarfile, io
    buf = io.BytesIO()
    with tarfile.open(fileobj=buf, mode="w:gz") as tar:
        for db_name in ["api_keys.json", "usage.json", "memory.db", "cache.db",
                        "literary.db", "scheduler.db", "webhooks.db", "versions.db"]:
            path = f"/data/{db_name}"
            if os.path.exists(path):
                tar.add(path, arcname=db_name)
    buf.seek(0)
    return Response(buf.read(), media_type="application/gzip",
                    headers={"Content-Disposition": "attachment; filename=intello_backup.tar.gz"})


# --- Rate Limit Dashboard (#8) ---

@app.get("/api/usage/history")
async def api_usage_history():
    """Get usage history across all providers."""
    from intello import ratelimit
    usage = ratelimit._load()
    result = {}
    for day, models in usage.items():
        result[day] = models
    # Also add current remaining
    current = {}
    for p in _providers:
        if p.available:
            rem = ratelimit.remaining(p.model_id, p.daily_limit)
            current[p.model_id] = {"name": p.name, "used": ratelimit.get_usage(p.model_id),
                                    "limit": p.daily_limit, "remaining": rem}
    return {"history": result, "today": current}


# --- Prompt Templates (#6) ---

PROMPT_TEMPLATES = {
    "analyze_pacing": {"name": "Analyze Pacing", "prompt": "Analyze the pacing of chapter {chapter}. Where is it too slow or fast?"},
    "character_check": {"name": "Character Consistency", "prompt": "Check {character} for consistency across all chapters. Flag any contradictions."},
    "show_not_tell": {"name": "Show Not Tell", "prompt": "Find all instances of telling instead of showing in chapter {chapter} and rewrite them."},
    "tighten_prose": {"name": "Tighten Prose", "prompt": "Tighten the prose in lines {start}-{end}. Remove unnecessary words, strengthen verbs."},
    "expand_scene": {"name": "Expand Scene", "prompt": "Expand the scene at lines {start}-{end} with more sensory detail and character interiority."},
    "blurb": {"name": "Generate Blurb", "prompt": "Write a compelling back-cover blurb for this book."},
    "chapter_summary": {"name": "Chapter Summary", "prompt": "Summarize each chapter in one sentence."},
}


@app.get("/api/templates")
async def api_templates():
    return PROMPT_TEMPLATES


# --- Scheduler background loop ---

async def _scheduler_loop():
    """Background task that runs due scheduled tasks every 60 seconds."""
    while True:
        await asyncio.sleep(60)
        try:
            due = scheduler.get_due_tasks()
            for task in due:
                plan = build_plan(task["prompt"], _providers)
                if plan.primary:
                    result = await execute(plan.primary, task["prompt"], max_tokens=2000)
                    scheduler.record_result(task["task_id"],
                                            result.content if not result.degraded else "Failed")
        except Exception:
            pass


@app.on_event("startup")
async def start_scheduler():
    asyncio.create_task(_scheduler_loop())


# --- OCR Service ---

@app.post("/api/v1/ocr")
async def api_ocr_image(
    file: UploadFile = File(...),
    language: str = Form("eng"),
    output: str = Form("json"),
    quality: str = Form("auto"),  # fast | auto | best
):
    """OCR a single image. quality=auto escalates engines on low confidence."""
    import tempfile
    with tempfile.NamedTemporaryFile(suffix=f"_{file.filename}", delete=False) as f:
        content = await file.read()
        f.write(content)
        tmp = f.name

    result = await ocr_engines.smart_ocr(tmp, language, quality)
    os.unlink(tmp)

    if output == "text":
        return Response(result["text"], media_type="text/plain")
    return result


@app.post("/api/v1/ocr/pdf")
async def api_ocr_pdf(
    file: UploadFile = File(...),
    language: str = Form("eng"),
    output: str = Form("json"),
    pages: str = Form(""),
):
    """OCR a PDF — returns text or searchable PDF."""
    import tempfile
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
        content = await file.read()
        f.write(content)
        tmp = f.name

    if output == "searchable_pdf":
        out_path = tmp + "_ocr.pdf"
        ok = ocr.ocr_pdf_searchable(tmp, out_path, language, pages)
        os.unlink(tmp)
        if ok:
            from fastapi.responses import FileResponse
            return FileResponse(out_path, media_type="application/pdf",
                                filename=f"ocr_{file.filename}")
        return {"error": "OCR failed"}

    result = ocr.ocr_pdf_to_text(tmp, language, pages)
    os.unlink(tmp)

    if output == "text":
        full_text = "\n\n".join(f"--- Page {p['page']} ---\n{p['text']}" for p in result["pages"])
        return Response(full_text, media_type="text/plain")
    return result


@app.post("/api/v1/ocr/jobs")
async def api_ocr_create_job(
    file: Optional[UploadFile] = File(None),
    file_url: Optional[str] = Form(None),
    language: str = Form("eng"),
    output: str = Form("searchable_pdf"),
    pages: str = Form(""),
):
    """Create an async OCR job for large PDFs."""
    import tempfile, httpx

    if file and file.filename:
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False, dir=str(ocr.JOBS_DIR)) as f:
            f.write(await file.read())
            tmp = f.name
    elif file_url:
        async with httpx.AsyncClient(timeout=120) as c:
            r = await c.get(file_url)
            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False, dir=str(ocr.JOBS_DIR)) as f:
                f.write(r.content)
                tmp = f.name
    else:
        return {"error": "Provide file or file_url"}

    job_id = ocr.create_job(tmp, language, output, pages)
    asyncio.create_task(ocr.run_job(job_id))
    return ocr.get_job(job_id)


@app.get("/api/v1/ocr/jobs/{job_id}")
async def api_ocr_job_status(job_id: str):
    job = ocr.get_job(job_id)
    if not job:
        return {"error": "Job not found"}
    return {k: v for k, v in job.items() if k != "file_path"}


@app.get("/api/v1/ocr/jobs/{job_id}/result")
async def api_ocr_job_result(job_id: str):
    job = ocr.get_job(job_id)
    if not job or job["status"] != "complete" or not job.get("result_path"):
        return {"error": "Job not complete"}

    if job["result_path"].endswith(".pdf"):
        from fastapi.responses import FileResponse
        return FileResponse(job["result_path"], media_type="application/pdf")
    else:
        with open(job["result_path"]) as f:
            return json.loads(f.read())


# --- OpenAI-compatible API (R2) ---

@app.post("/v1/chat/completions")
async def openai_chat_completions(request: Request):
    """OpenAI-compatible chat/completions endpoint. Supports stream:true, proper errors, timeouts."""
    user = _get_user(request)
    user_provs = filter_providers_for_user(_providers, user)
    body = await request.json()
    messages = body.get("messages", [])
    max_tokens = body.get("max_tokens", 4096)
    model_hint = body.get("model", "")
    prefer_free = body.get("prefer_free", True)
    stream = body.get("stream", False)

    # Extract system + user messages
    system_msg = None
    user_msg = ""
    for m in messages:
        if m["role"] == "system":
            system_msg = m["content"]
        elif m["role"] == "user":
            user_msg = m["content"]

    if not user_msg:
        return JSONResponse({"error": {"message": "No user message", "type": "invalid_request"}}, 400)

    # Route
    plan = build_plan(user_msg, user_provs, prefer_free=prefer_free)

    if model_hint:
        for p in user_provs:
            if p.available and (model_hint in p.model_id or model_hint in p.name.lower()):
                plan.primary = p
                break

    # Fix #4: Return 429 when all providers exhausted
    if not plan.primary:
        all_exhausted = all(
            ratelimit.remaining(p.model_id, p.daily_limit) == 0
            for p in user_provs if p.available
        )
        if all_exhausted:
            return JSONResponse(
                {"error": {"message": "All providers rate-limited. Try again later.", "type": "rate_limit_exhausted"}},
                status_code=429,
                headers={"Retry-After": "3600"}
            )
        return JSONResponse(
            {"error": {"message": "No providers available", "type": "server_error",
                       "missing_keys": plan.missing_keys}},
            status_code=503
        )

    # Fix #5: Cache key includes system prompt
    cache_key = f"{system_msg or ''}|||{user_msg}"
    cached = cache.get_cached(cache_key, plan.task_type.value)
    if cached and not stream:
        return _openai_response(cached["response"], cached["provider"], cached["model"], 0, 0, True)

    # Fix #3: Handle stream:true in the main endpoint
    if stream:
        return await _stream_response(user_msg, system_msg, max_tokens, plan, user_provs)

    # Execute with fallback chain + Fix #1 (timeout) + Fix #2 (structured errors)
    chain = [plan.primary] + plan.fallbacks
    last_error = ""
    providers_tried = []
    for provider in chain:
        if not provider or not provider.available:
            continue
        providers_tried.append(provider.name)
        result = await execute(provider, user_msg, max_tokens=max_tokens, system=system_msg)
        if not result.degraded:
            cache.store(cache_key, plan.task_type.value, result.content,
                        result.provider_name, result.model_id, result.cost)
            resp = _openai_response(result.content, result.provider_name, result.model_id,
                                    result.input_tokens, result.output_tokens, False)
            resp["x_intello"]["providers_tried"] = providers_tried
            resp["x_intello"]["fallback_count"] = len(providers_tried) - 1
            return resp
        last_error = result.content

    # Fix #2: Structured error with provider info
    return JSONResponse({
        "error": {
            "message": f"All providers failed. Last error: {last_error}",
            "type": "provider_error",
            "providers_tried": providers_tried,
            "fallback_count": len(providers_tried),
        }
    }, status_code=502)


async def _stream_response(user_msg, system_msg, max_tokens, plan, providers):
    """SSE streaming for /v1/chat/completions with stream:true."""
    async def generate():
        provider = plan.primary
        try:
            from openai import AsyncOpenAI
            base_urls = {"openai": None, "groq": "https://api.groq.com/openai/v1",
                         "mistral": "https://api.mistral.ai/v1", "deepseek": "https://api.deepseek.com",
                         "openrouter": "https://openrouter.ai/api/v1", "xai": "https://api.x.ai/v1"}
            base = base_urls.get(provider.provider)
            if base is not None or provider.provider == "openai":
                kwargs = {"api_key": provider.api_key}
                if base:
                    kwargs["base_url"] = base
                client = AsyncOpenAI(**kwargs)
                msgs = [{"role": "system", "content": system_msg or "You are a helpful assistant."},
                        {"role": "user", "content": user_msg}]
                stream = await asyncio.wait_for(
                    client.chat.completions.create(
                        model=provider.model_id, messages=msgs, max_tokens=max_tokens, stream=True),
                    timeout=30
                )
                async for chunk in stream:
                    if chunk.choices and chunk.choices[0].delta.content:
                        text = chunk.choices[0].delta.content
                        yield f"data: {json.dumps({'choices': [{'delta': {'content': text}, 'index': 0}]})}\n\n"
                yield f"data: {json.dumps({'choices': [{'delta': {}, 'finish_reason': 'stop', 'index': 0}], 'x_intello': {'provider': provider.name}})}\n\n"
                yield "data: [DONE]\n\n"
                return
        except Exception:
            pass
        # Fallback: non-streaming, send all at once
        result = await execute(provider, user_msg, max_tokens=max_tokens, system=system_msg)
        yield f"data: {json.dumps({'choices': [{'delta': {'content': result.content}, 'index': 0}]})}\n\n"
        yield f"data: {json.dumps({'choices': [{'delta': {}, 'finish_reason': 'stop', 'index': 0}], 'x_intello': {'provider': result.provider_name}})}\n\n"
        yield "data: [DONE]\n\n"

    return StreamingResponse(generate(), media_type="text/event-stream")



@app.post("/v1/chat/completions/stream")
async def openai_chat_stream(request: Request):
    """Legacy streaming endpoint — redirects to main endpoint with stream:true."""
    body = await request.json()
    body["stream"] = True
    # Reconstruct request with stream flag
    from starlette.requests import Request as StarletteRequest
    return await openai_chat_completions(request)

def _openai_response(content, provider, model, inp_tokens, out_tokens, was_cached):
    """Format response in OpenAI chat/completions format."""
    import time as _time
    return {
        "id": f"chatcmpl-{uuid.uuid4().hex[:12]}",
        "object": "chat.completion",
        "created": int(_time.time()),
        "model": model,
        "choices": [{
            "index": 0,
            "message": {"role": "assistant", "content": content},
            "finish_reason": "stop",
        }],
        "usage": {
            "prompt_tokens": inp_tokens,
            "completion_tokens": out_tokens,
            "total_tokens": inp_tokens + out_tokens,
        },
        "x_intello": {
            "provider": provider,
            "cached": was_cached,
        },
    }


# --- Status endpoint (R5) ---

@app.get("/api/v1/status")
async def api_status():
    avail = [p for p in _providers if p.available]
    free = [p for p in avail if p.tier == Tier.FREE]
    return {
        "available": len(avail) > 0,
        "providers": [
            {"name": p.name, "model": p.model_id, "tier": p.tier.value,
             "available": p.available, "provider": p.provider}
            for p in _providers
        ],
        "total_available": len(avail),
        "free_available": len(free),
        "ocr": {
            "available": shutil.which("tesseract") is not None,
            "engines": [
                {"name": "tesseract", "type": "local", "available": shutil.which("tesseract") is not None},
                {"name": "ocr.space", "type": "cloud_free", "available": True},
                {"name": "gemini_vision", "type": "llm", "available": any(
                    p.available and p.provider == "google" for p in _providers)},
            ],
            "languages": ocr.get_languages(),
            "quality_modes": ["fast", "auto", "best"],
        },
        "speech": {
            "tts_available": speech.tts_available(),
            "tts_engine": "piper",
            "tts_voices": [v["id"] for v in speech.get_available_voices()],
            "stt_provider": "groq",
            "stt_model": "whisper-large-v3-turbo",
            "stt_daily_limit_seconds": 28800,
        },
    }


# --- Models list (OpenAI-compatible) ---

@app.get("/v1/models")
async def openai_models():
    return {
        "object": "list",
        "data": [
            {"id": p.model_id, "object": "model", "owned_by": p.provider,
             "created": 1700000000}
            for p in _providers if p.available
        ],
    }
